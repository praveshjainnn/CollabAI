import io
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches

class TiptapHTMLParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        
        # Formatting states
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.is_heading = False
        self.heading_level = 0
        self.is_list = False
        self.list_style = None  # 'bullet' or 'numbered'
        self.list_item = False
        self.blockquote = False

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.is_heading = True
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_heading(level=self.heading_level)
            
        elif tag == 'p':
            if not self.list_item:
                if self.blockquote:
                    self.current_paragraph = self.doc.add_paragraph(style='Quote')
                else:
                    self.current_paragraph = self.doc.add_paragraph()
                
        elif tag == 'blockquote':
            self.blockquote = True
            
        elif tag == 'ul':
            self.is_list = True
            self.list_style = 'List Bullet'
            
        elif tag == 'ol':
            self.is_list = True
            self.list_style = 'List Number'
            
        elif tag == 'li':
            self.list_item = True
            self.current_paragraph = self.doc.add_paragraph(style=self.list_style or 'List Bullet')
            
        elif tag == 'strong' or tag == 'b':
            self.is_bold = True
            
        elif tag == 'em' or tag == 'i':
            self.is_italic = True
            
        elif tag == 'u':
            self.is_underline = True

    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.is_heading = False
            self.heading_level = 0
            self.current_paragraph = None
            self.current_run = None
            
        elif tag == 'p':
            if not self.list_item:
                self.current_paragraph = None
                self.current_run = None
                
        elif tag == 'blockquote':
            self.blockquote = False
            
        elif tag in ['ul', 'ol']:
            self.is_list = False
            self.list_style = None
            
        elif tag == 'li':
            self.list_item = False
            self.current_paragraph = None
            self.current_run = None
            
        elif tag == 'strong' or tag == 'b':
            self.is_bold = False
            
        elif tag == 'em' or tag == 'i':
            self.is_italic = False
            
        elif tag == 'u':
            self.is_underline = False

    def handle_data(self, data):
        if not data.strip() and not data == " ":
            return
            
        # Ensure we have a paragraph to add to
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
            
        self.current_run = self.current_paragraph.add_run(data)
        
        # Apply formatting states
        if self.is_bold:
            self.current_run.bold = True
        if self.is_italic:
            self.current_run.italic = True
        if self.is_underline:
            self.current_run.underline = True

def convert_html_to_docx(html_content: str, title: str) -> bytes:
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title formatting
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_p.paragraph_format.space_after = Pt(20)
    
    # Parse HTML and populate docx
    parser = TiptapHTMLParser(doc)
    parser.feed(html_content)
    
    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()
